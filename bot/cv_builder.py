from telegram import Update, ReplyKeyboardRemove
from telegram.ext import (
    ContextTypes, ConversationHandler, CommandHandler, MessageHandler, filters
)
from docx import Document
from docx.shared import Pt
import os
from asgiref.sync import sync_to_async
from bot.models import User

# Extended conversation states
(
    NAME, TITLE, EMAIL, PHONE, LOCATION, LINKS, SUMMARY,
    EDUCATION, ADD_EDUCATION, EXPERIENCE, ADD_EXPERIENCE,
    CERTIFICATIONS, ADD_CERTIFICATIONS,
    LANGUAGES, AWARDS, ADD_AWARDS,
    REFEREES, ADD_REFEREES, SKILLS
) = range(19)


async def start_cv(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message is None:
        return ConversationHandler.END

    context.user_data.clear()
    context.user_data['education'] = []
    context.user_data['experience'] = []
    context.user_data['certifications'] = []
    context.user_data['awards'] = []
    context.user_data['referees'] = []

    await update.message.reply_text(
        "👋 *Let's build your professional CV!* \n\n"
        "I'll ask you a few questions to create a stunning resume.\n"
        "First, what is your *Full Name*?",
        parse_mode='Markdown'
    )
    return NAME


async def get_name(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['name'] = update.message.text
    await update.message.reply_text(
        "Great! What is your *Current Professional Title*? (e.g., Software Engineer, Marketing Manager)",
        parse_mode='Markdown'
    )
    return TITLE


async def get_title(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['title'] = update.message.text
    await update.message.reply_text("📧 What is your *Email Address*?", parse_mode='Markdown')
    return EMAIL


async def get_email(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['email'] = update.message.text
    await update.message.reply_text("📱 What is your *Phone Number*?", parse_mode='Markdown')
    return PHONE


async def get_phone(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['phone'] = update.message.text
    await update.message.reply_text("🌍 What is your *City and Country* of residence?", parse_mode='Markdown')
    return LOCATION


async def get_location(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['location'] = update.message.text
    await update.message.reply_text(
        "🔗 Please provide your *Professional Links* (LinkedIn, Portfolio, etc.), separated by commas:",
        parse_mode='Markdown'
    )
    return LINKS


async def get_links(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['links'] = [l.strip() for l in update.message.text.split(',')]
    await update.message.reply_text(
        "📝 Write a *Brief Professional Summary* (2-3 sentences about your experience and goals):",
        parse_mode='Markdown'
    )
    return SUMMARY


async def get_summary(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['summary'] = update.message.text
    await update.message.reply_text(
        "🎓 Let's add your *Education*.\n"
        "Please enter an entry in this format:\n"
        "_Degree, University, Year_",
        parse_mode='Markdown'
    )
    return EDUCATION


async def get_education(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['education'].append(update.message.text)
    await update.message.reply_text("Do you want to add another education entry? (yes/no)")
    return ADD_EDUCATION


async def add_education(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == 'yes':
        await update.message.reply_text("Enter the next education entry:")
        return EDUCATION
    await update.message.reply_text(
        "💼 Now, let's add your *Work Experience*.\n"
        "Please enter your most recent role:\n"
        "_Role, Company, Duration (e.g. 2020-Present)_",
        parse_mode='Markdown'
    )
    return EXPERIENCE


async def get_experience(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['experience'].append(update.message.text)
    await update.message.reply_text("Add another work experience entry? (yes/no)")
    return ADD_EXPERIENCE


async def add_experience(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == 'yes':
        await update.message.reply_text("Enter the next work experience:")
        return EXPERIENCE
    await update.message.reply_text(
        "📜 Do you have any *Certifications*? (e.g., PMP, AWS Certified)\n"
        "Enter one, or type 'skip' if none.",
        parse_mode='Markdown'
    )
    return CERTIFICATIONS


async def get_certifications(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    text = update.message.text
    if text.lower() != 'skip':
        context.user_data['certifications'].append(text)
        await update.message.reply_text("Add another certification? (yes/no)")
        return ADD_CERTIFICATIONS
    
    await update.message.reply_text("🗣️ What *Languages* do you speak? (comma-separated)", parse_mode='Markdown')
    return LANGUAGES


async def add_certifications(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == 'yes':
        await update.message.reply_text("Enter the next certification:")
        return CERTIFICATIONS
    await update.message.reply_text("🗣️ What *Languages* do you speak? (comma-separated)", parse_mode='Markdown')
    return LANGUAGES


async def get_languages(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['languages'] = [l.strip() for l in update.message.text.split(',')]
    await update.message.reply_text(
        "🏆 Any *Awards or Achievements*?\n"
        "Enter one, or type 'skip' if none.",
        parse_mode='Markdown'
    )
    return AWARDS


async def get_awards(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    text = update.message.text
    if text.lower() != 'skip':
        context.user_data['awards'].append(text)
        await update.message.reply_text("Add another award? (yes/no)")
        return ADD_AWARDS
    
    await update.message.reply_text(
        "👥 Please add a *Referee*:\n"
        "_Name, Position, Contact Info_",
        parse_mode='Markdown'
    )
    return REFEREES


async def add_awards(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == 'yes':
        await update.message.reply_text("Enter the next award:")
        return AWARDS
    await update.message.reply_text(
        "👥 Please add a *Referee*:\n"
        "_Name, Position, Contact Info_",
        parse_mode='Markdown'
    )
    return REFEREES


async def get_referees(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    context.user_data['referees'].append(update.message.text)
    await update.message.reply_text("Add another referee? (yes/no)")
    return ADD_REFEREES


async def add_referees(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message.text.lower() == 'yes':
        await update.message.reply_text("Enter the next referee:")
        return REFEREES
    await update.message.reply_text(
        "💡 Finally, list your *Key Skills* (comma-separated):\n"
        "_e.g., Python, Project Management, SEO_",
        parse_mode='Markdown'
    )
    return SKILLS


async def get_skills(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    d = context.user_data
    d['skills'] = [s.strip() for s in update.message.text.split(',')]

    user = await sync_to_async(User.objects.get)(user_id=str(update.effective_user.id))
    user.cv_data = d
    user.current_job_title = d['title']
    user.skills = d['skills']
    await sync_to_async(user.save)()

    await sync_to_async(user.save)()
    
    await update.message.reply_text("✨ *Generating your professional CV...* Please wait a moment.", parse_mode='Markdown')

    # Create CV Document
    doc = Document()
    
    # --- STYLE SETTINGS FOR ATS ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Helper to add cleaner sections
    def add_section_header(document, text):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Arial'
        # Add a bottom border if possible, but for simplicity we'll just use spacing
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    # Header
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(d['name'])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.name = 'Arial'
    name_p.alignment = 1 # Center

    info_p = doc.add_paragraph()
    info_p.alignment = 1
    info_run = info_p.add_run(f"{d['title']}\n{d['email']} | {d['phone']} | {d['location']}")
    info_run.font.size = Pt(10)

    if d.get('links'):
        links_p = doc.add_paragraph()
        links_p.alignment = 1
        links_run = links_p.add_run(" | ".join(d['links']))
        links_run.font.size = Pt(10)

    doc.add_paragraph() # Spacer

    # Professional Summary
    add_section_header(doc, "PROFESSIONAL SUMMARY")
    doc.add_paragraph(d['summary'])

    # Experience
    add_section_header(doc, "PROFESSIONAL EXPERIENCE")
    for exp in d['experience']:
        p = doc.add_paragraph(exp, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Education
    add_section_header(doc, "EDUCATION")
    for edu in d['education']:
        p = doc.add_paragraph(edu, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

    # Skills
    add_section_header(doc, "SKILLS")
    skills_text = ", ".join(d['skills'])
    doc.add_paragraph(skills_text)

    # Certifications
    if d['certifications']:
        add_section_header(doc, "CERTIFICATIONS")
        for cert in d['certifications']:
            doc.add_paragraph(cert, style='List Bullet')

    # Languages
    if d['languages']:
        add_section_header(doc, "LANGUAGES")
        doc.add_paragraph(", ".join(d['languages']))

    # Awards
    if d['awards']:
        add_section_header(doc, "AWARDS")
        for award in d['awards']:
            doc.add_paragraph(award, style='List Bullet')

    # Referees
    if d['referees']:
        add_section_header(doc, "REFERENCES")
        for ref in d['referees']:
            doc.add_paragraph(ref, style='List Bullet')

    filename = f"{d['name'].replace(' ', '_')}_CV.docx"
    doc.save(filename)
    with open(filename, 'rb') as f:
        await update.message.reply_document(
            document=f,
            caption="✅ *Here is your new CV!* \nGood luck with your job search! 🚀",
            parse_mode='Markdown'
        )
    os.remove(filename)

    return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    if update.message is not None:
        await update.message.reply_text("CV building canceled.", reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END


def get_cv_handler():
    return ConversationHandler(
        entry_points=[
            CommandHandler('build_cv', start_cv),
            CommandHandler('cv', start_cv),
        ],
        states={
            NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_name)],
            TITLE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_title)],
            EMAIL: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_email)],
            PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_phone)],
            LOCATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_location)],
            LINKS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_links)],
            SUMMARY: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_summary)],
            EDUCATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_education)],
            ADD_EDUCATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_education)],
            EXPERIENCE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_experience)],
            ADD_EXPERIENCE: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_experience)],
            CERTIFICATIONS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_certifications)],
            ADD_CERTIFICATIONS: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_certifications)],
            LANGUAGES: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_languages)],
            AWARDS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_awards)],
            ADD_AWARDS: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_awards)],
            REFEREES: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_referees)],
            ADD_REFEREES: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_referees)],
            SKILLS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_skills)],
        },
        fallbacks=[CommandHandler('cancel', cancel)]
    )